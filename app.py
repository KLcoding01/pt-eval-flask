import os
import io
import re
import json
from sqlalchemy import or_
from flask import Flask, request, jsonify, redirect, url_for, flash, render_template, send_file, session, abort, render_template_string
from flask_mail import Mail, Message
from flask_login import LoginManager, login_user, logout_user, login_required, current_user, UserMixin
from werkzeug.security import generate_password_hash, check_password_hash
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from datetime import date, datetime, timedelta
from io import BytesIO
from functools import wraps
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import Column, Integer, String, Boolean, DateTime, ForeignKey, func
from sqlalchemy.dialects.sqlite import JSON
from flask_migrate import Migrate
from itsdangerous import URLSafeTimedSerializer
from dateutil import parser

# Google Calendar imports
from google_auth_oauthlib.flow import Flow
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from google.auth.transport.requests import Request

# DB MODELS
from models import db, CPTCode, ICD10Code, Patient, Visit, Attachment, Billing, Therapist, Physician, Insurance, PTNote

# ====== ENV & CONFIG ======
load_dotenv()
app = Flask(__name__)

app.secret_key = os.getenv("SECRET_KEY", "dev_secret_key_change_me")
serializer = URLSafeTimedSerializer(app.secret_key)

basedir = os.path.abspath(os.path.dirname(__file__))
db_path = os.path.join(basedir, 'static', 'uploads', 'db.sqlite3')
app.config['SQLALCHEMY_DATABASE_URI'] = f"sqlite:///{db_path}"
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'static/uploads'

# ====== FORGOT PASSWORD RESET ======
app.config['MAIL_SERVER'] = os.getenv("MAIL_SERVER", "smtp.gmail.com")
app.config['MAIL_PORT'] = int(os.getenv("MAIL_PORT", 587))
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.getenv("MAIL_USERNAME", "your_email@gmail.com")
app.config['MAIL_PASSWORD'] = os.getenv("MAIL_PASSWORD", "your_app_password")
app.config['MAIL_DEFAULT_SENDER'] = os.getenv("MAIL_DEFAULT_SENDER", "your_email@gmail.com")
mail = Mail(app)

#======= OPEN AI=======
OPENAI_KEY = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=OPENAI_KEY) if OPENAI_KEY else None
MODEL = "gpt-4o-mini"

# ====== INIT DB ======
db.init_app(app)
migrate = Migrate(app, db)

# ====== LOGIN MANAGER ======
login_manager = LoginManager()
login_manager.login_view = 'login'
login_manager.init_app(app)

@login_manager.user_loader
def load_user(user_id):
    return Therapist.query.get(int(user_id))

# ====== CREATE TABLES & ADMIN IF NONE ======
with app.app_context():
    db.create_all()
    # Only run this once! Add a default admin user if none exists:
    if not Therapist.query.filter_by(username="admin").first():
        hashed_pw = generate_password_hash("admin123")
        admin = Therapist(
            username="admin",
            password=hashed_pw,
            first_name="Admin",
            last_name="User",
            credentials="PT",
            email="admin@example.com",
            phone="555-555-5555",
            availability="M-F",
            npi="0000000000",
            pt_license="LICENSE123",
        )
        db.session.add(admin)
        db.session.commit()
        print("Default admin user created: username=admin, password=admin123")

def parse_dob(dob_str):
    """
    Parse DOB string in MM-DD-YYYY or MM/DD/YYYY format.
    Returns a date object or None if invalid.
    """
    if not dob_str:
        return None
    for fmt in ("%m-%d-%Y", "%m/%d/%Y"):
        try:
            return datetime.strptime(dob_str, fmt).date()
        except ValueError:
            continue
    return None

# ====== ROUTES ======

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = Therapist.query.filter_by(username=username).first()
        if user and check_password_hash(user.password, password):
            login_user(user)
            return redirect(url_for('dashboard'))
        else:
            flash("Invalid username or password", "danger")
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

# ... (forgot-password, reset-password, register, dashboard routes) ...

@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html')

# ====== GOOGLE CALENDAR INTEGRATION ======

SCOPES = ['https://www.googleapis.com/auth/calendar']

def get_google_calendar_service():
    creds_data = session.get('credentials')
    if not creds_data:
        return None
    creds = Credentials(**creds_data)
    if creds.expired and creds.refresh_token:
        creds.refresh(Request())
        session['credentials'] = {
            'token': creds.token,
            'refresh_token': creds.refresh_token,
            'token_uri': creds.token_uri,
            'client_id': creds.client_id,
            'client_secret': creds.client_secret,
            'scopes': creds.scopes
        }
    service = build('calendar', 'v3', credentials=creds)
    return service

def create_google_event(visit):
    service = get_google_calendar_service()
    if not service:
        return None
    event_body = {
        'summary': f"{visit.patient.first_name} {visit.patient.last_name} - {visit.visit_type}",
        'start': {'dateTime': visit.visit_date.isoformat(), 'timeZone': 'America/Los_Angeles'},
        'end': {'dateTime': (visit.visit_date + timedelta(minutes=getattr(visit, 'duration', 60))).isoformat(), 'timeZone': 'America/Los_Angeles'},
        'description': f"Therapist ID: {visit.therapist_id}"
    }
    event = service.events().insert(calendarId='primary', body=event_body).execute()
    return event.get('id')

def update_google_event(visit):
    service = get_google_calendar_service()
    if not service or not visit.google_event_id:
        return False
    event_body = {
        'summary': f"{visit.patient.first_name} {visit.patient.last_name} - {visit.visit_type}",
        'start': {'dateTime': visit.visit_date.isoformat(), 'timeZone': 'America/Los_Angeles'},
        'end': {'dateTime': (visit.visit_date + timedelta(minutes=getattr(visit, 'duration', 60))).isoformat(), 'timeZone': 'America/Los_Angeles'},
        'description': f"Therapist ID: {visit.therapist_id}"
    }
    service.events().update(calendarId='primary', eventId=visit.google_event_id, body=event_body).execute()
    return True

def delete_google_event(visit):
    service = get_google_calendar_service()
    if not service or not visit.google_event_id:
        return False
    try:
        service.events().delete(calendarId='primary', eventId=visit.google_event_id).execute()
        return True
    except Exception as e:
        print(f"Google Calendar delete event error: {e}")
        return False

@app.route('/authorize')
@login_required
def authorize():
    flow = Flow.from_client_secrets_file(
        'credentials.json',
        scopes=SCOPES,
        redirect_uri=url_for('oauth2callback', _external=True)
    )
    authorization_url, state = flow.authorization_url(
        access_type='offline',
        include_granted_scopes='true'
    )
    session['state'] = state
    return redirect(authorization_url)

@app.route('/oauth2callback')
@login_required
def oauth2callback():
    state = session.get('state')
    flow = Flow.from_client_secrets_file(
        'credentials.json',
        scopes=SCOPES,
        state=state,
        redirect_uri=url_for('oauth2callback', _external=True)
    )
    flow.fetch_token(authorization_response=request.url)
    credentials = flow.credentials
    session['credentials'] = {
        'token': credentials.token,
        'refresh_token': credentials.refresh_token,
        'token_uri': credentials.token_uri,
        'client_id': credentials.client_id,
        'client_secret': credentials.client_secret,
        'scopes': credentials.scopes
    }
    flash("Google Calendar connected!", "success")
    return redirect(url_for('dashboard'))

# Core API routes for events CRUD, with timezone-aware parsing

@app.route('/api/events', methods=['GET', 'POST'])
@login_required
def api_events():
    if request.method == 'GET':
        therapist_id = request.args.get('therapist_id', type=int)
        if therapist_id:
            visits = Visit.query.filter_by(therapist_id=therapist_id).all()
        else:
            visits = Visit.query.all()
        app.logger.info(f"Fetching events for therapist_id={therapist_id}, found {len(visits)} visits")
        events = []
        for visit in visits:
            events.append({
                "id": visit.id,
                "title": visit.visit_type or "Untitled",
                "start": visit.visit_date.isoformat(),
                "end": (visit.visit_date + timedelta(minutes=visit.duration)).isoformat() if visit.duration else None,
                "extendedProps": {
                    "therapist_id": visit.therapist_id,
                    "patient_id": visit.patient_id
                }
            })
        return jsonify(events)

    elif request.method == 'POST':
        data = request.json
        therapist_id = data.get('therapist_id')
        patient_id = data.get('patient_id')
        title = data.get('title')
        start = data.get('start')
        end = data.get('end')

        if not all([therapist_id, patient_id, title, start, end]):
            return jsonify({'error': 'Missing required fields'}), 400

        try:
            visit = Visit(
                patient_id=patient_id,
                therapist_id=therapist_id,
                visit_date=parser.parse(start),  # dateutil parse ISO8601 with tz
                duration=int((parser.parse(end) - parser.parse(start)).total_seconds() / 60),
                visit_type=title,
                status='Scheduled'
            )
            db.session.add(visit)
            db.session.commit()

            event_id = create_google_event(visit)
            if event_id:
                visit.google_event_id = event_id
                db.session.commit()

            return jsonify({
                'success': True,
                'id': visit.id,
                'google_event_id': visit.google_event_id,
            })

        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500


@app.route('/api/events/<int:event_id>', methods=['PUT'])
@login_required
def api_update_event(event_id):
    visit = Visit.query.get_or_404(event_id)
    data = request.json
    therapist_id = data.get('therapist_id')
    patient_id = data.get('patient_id')
    title = data.get('title')
    start = data.get('start')
    end = data.get('end')

    if not all([therapist_id, patient_id, title, start, end]):
        return jsonify({'error': 'Missing required fields'}), 400

    try:
        visit.therapist_id = therapist_id
        visit.patient_id = patient_id
        visit.visit_type = title
        visit.visit_date = datetime.fromisoformat(start)
        visit.duration = int((datetime.fromisoformat(end) - datetime.fromisoformat(start)).total_seconds() / 60)
        db.session.commit()

        if visit.google_event_id:
            update_google_event(visit)
        else:
            event_id = create_google_event(visit)
            if event_id:
                visit.google_event_id = event_id
                db.session.commit()

        return jsonify({'success': True})

    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/events/<int:event_id>', methods=['DELETE'])
@login_required
def api_delete_event(event_id):
    visit = Visit.query.get_or_404(event_id)

    try:
        if visit.google_event_id:
            delete_google_event(visit)
        db.session.delete(visit)
        db.session.commit()
        return jsonify({'success': True})

    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

# Therapist list for dropdown
@app.route('/api/therapist_list')
@login_required
def therapist_list():
    therapists = Therapist.query.all()
    return jsonify([{'id': t.id, 'first_name': t.first_name, 'last_name': t.last_name} for t in therapists])

# Patient list for dropdown
@app.route('/api/patient_list')
@login_required
def api_patient_list():
    patients = Patient.query.order_by(Patient.last_name).all()
    return jsonify([
        {
            "id": p.id,
            "first_name": p.first_name,
            "last_name": p.last_name,
            "dob": p.dob.strftime('%m-%d-%Y') if p.dob else ""
        }
        for p in patients
    ])

@app.route('/pt_save_to_patient', methods=['POST'])
@login_required
def pt_save_to_patient():
    data = request.get_json()
    therapist_id = data.get('therapist_id')
    patient_id = data.get('patient_id')
    visit_type = data.get('visit_type', 'PT Evaluation')

    if not therapist_id or not patient_id:
        return jsonify({'status': 'error', 'message': 'Therapist and Patient required.'}), 400

    # Copy data and remove non-note fields
    notes_dict = dict(data)
    for k in ['therapist_id', 'patient_id', 'visit_type', 'duration']:
        notes_dict.pop(k, None)

    # Trim whitespace/newlines for all string values in notes_dict
    for key, val in notes_dict.items():
        if isinstance(val, str):
            notes_dict[key] = val.strip()

    visit = Visit(
        patient_id=patient_id,
        therapist_id=therapist_id,
        visit_date=datetime.utcnow(),
        visit_type=visit_type,
        status='Completed',
        notes=json.dumps(notes_dict),
    )
    db.session.add(visit)
    db.session.commit()
    return jsonify({'status': 'ok', 'message': 'Visit saved!', 'visit_id': visit.id})
    
@app.template_filter('fromjson')
def fromjson_filter(s):
    try:
        return json.loads(s)
    except Exception:
        return {}
        
@app.route('/edit_visit_date/<int:visit_id>', methods=['POST'])
def edit_visit_date(visit_id):
    new_date = request.form.get('date')
    visit = Visit.query.get_or_404(visit_id)
    visit.visit_date = datetime.strptime(new_date, '%m-%d-%Y')
    db.session.commit()
    return redirect(url_for('patient_profile', patient_id=visit.patient_id))

@app.route("/visits")
@login_required
def visit_list():
    visits = Visit.query.order_by(Visit.visit_date.desc()).all()
    notes = PTNote.query.order_by(PTNote.date_created.desc()).all() 
    return render_template("visit_list.html", visits=visits, notes=notes)

@app.route("/visit/<int:visit_id>")
@login_required
def view_visit_note(visit_id):
    visit = Visit.query.get_or_404(visit_id)
    return render_template("view_visit_note.html", visit=visit)

@app.route("/patient/<int:patient_id>/notes")
@login_required
def patient_notes(patient_id):
    patient = Patient.query.get_or_404(patient_id)
    # Corrected this line to use date_created instead of created_at
    notes = PTNote.query.filter_by(patient_id=patient_id).order_by(PTNote.date_created.desc()).all()
    visits = Visit.query.filter_by(patient_id=patient_id).order_by(Visit.visit_date.desc()).all()
    return render_template("patient_notes.html", patient=patient, notes=notes, visits=visits)

    
@app.route('/add_visit', methods=['GET', 'POST'])
@login_required
def add_visit():
    if request.method == 'POST':
        patient_id = request.form.get('patient_id')
        therapist_id = request.form.get('therapist_id')
        visit_date_str = request.form.get('visit_date')
        end_time_str = request.form.get('end_time')
        duration = request.form.get('duration')
        visit_type = request.form.get('visit_type')
        status = request.form.get('status')
        cpt_code_id = request.form.get('cpt_code_id')
        icd10_code_id = request.form.get('icd10_code_id')
        notes = request.form.get('notes')
        google_event_id = request.form.get('google_event_id')

        # ---- FIX 1: Parse datetime ----
        visit_date = datetime.strptime(visit_date_str, "%m-%d-%Y %H:%M") if visit_date_str else None
        end_time = datetime.strptime(end_time_str, "%m-%d-%Y %H:%M") if end_time_str else None

        # ---- FIX 2: Validate patient_id ----
        if not patient_id:
            flash("Patient is required!", "danger")
            return redirect(request.url)
        
        # You may want to validate therapist_id and others too.

        visit = Visit(
            patient_id=patient_id,
            therapist_id=therapist_id,
            visit_date=visit_date,
            end_time=end_time,
            duration=duration,
            visit_type=visit_type,
            status=status,
            cpt_code_id=cpt_code_id if cpt_code_id else None,
            icd10_code_id=icd10_code_id if icd10_code_id else None,
            notes=notes,
            google_event_id=google_event_id
        )
        db.session.add(visit)
        db.session.commit()
        flash("Visit added!", "success")
        return redirect(url_for('visit_list'))

    # GET: show form
    patients = Patient.query.all()
    therapists = Therapist.query.all()
    cpt_codes = CPTCode.query.all()
    icd10_codes = ICD10Code.query.all()
    return render_template(
        'visit_form.html',
        patients=patients,
        therapists=therapists,
        cpt_codes=cpt_codes,
        icd10_codes=icd10_codes
    )


# ========== PATIENT CRUD ==========
def serialize_patient(patient):
    return {
        "id": patient.id,
        "first_name": patient.first_name,
        "last_name": patient.last_name,
        "mrn": patient.mrn,
        "dob": patient.dob.strftime("%m-%d-%Y") if patient.dob else None,
        "phone": patient.phone,
        "email": patient.email,
        "address": patient.address,
        "city": patient.city,
        "state": patient.state,
        "zip_code": patient.zip_code,
        "insurance_id": patient.insurance_id,
        "physician_id": patient.physician_id,
        "other_notes": other_notes
    }
    
@app.route('/api/patient/<int:patient_id>')
@login_required
def api_patient_detail(patient_id):
    patient = Patient.query.get_or_404(patient_id)
    return jsonify(serialize_patient(patient))
    
@app.route('/patients')
@login_required
def patients_list():
    patients = Patient.query.all()
    return render_template('patients.html', patients=patients, active_page='pt_builder')
    
@app.route('/patients/<int:patient_id>')
@login_required
def patient_detail(patient_id):
    patient = Patient.query.get_or_404(patient_id)
    deleted_cutoff = datetime.utcnow() - timedelta(days=30)

    # Fetch all visits
    all_visits = Visit.query.filter_by(patient_id=patient_id).order_by(Visit.visit_date.desc()).all()
    attachments = Attachment.query.filter_by(patient_id=patient_id).order_by(Attachment.uploaded_at.desc()).all()

    # Split visits into active and deleted
    deleted_visits = [v for v in all_visits if v.status == "Deleted"]
    for visit in deleted_visits:
        if visit.deleted_at and visit.deleted_at < deleted_cutoff:
            db.session.delete(visit)
    db.session.commit()

    active_visits = [v for v in all_visits if v.status != "Deleted"]

    # Active notes and deleted notes unchanged (optional)

    notes = PTNote.query.filter_by(patient_id=patient_id, deleted=False).order_by(PTNote.date_created.desc()).all()
    deleted_notes = PTNote.query.filter(
        PTNote.patient_id == patient_id,
        PTNote.deleted == True,
        or_(
            PTNote.deleted_at == None,
            PTNote.deleted_at >= deleted_cutoff
        )
    ).order_by(PTNote.deleted_at.desc()).all()

    edit_visit_id = request.args.get('edit_visit_id', type=int)

    return render_template(
        'patient_detail.html',
        patient=patient,
        visits=active_visits,
        deleted_visits=deleted_visits,
        attachments=attachments,
        notes=notes,
        deleted_notes=deleted_notes,
        edit_visit_id=edit_visit_id
        )
    
@app.route('/notes/<int:note_id>/delete', methods=['POST'])
@login_required
def delete_note(note_id):
    note = PTNote.query.get_or_404(note_id)
    try:
        note.deleted = True
        note.deleted_at = datetime.utcnow()
        db.session.commit()
        flash("Note deleted successfully.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error deleting note: {e}", "danger")
    return redirect(url_for('patient_detail', patient_id=note.patient_id))
    
@app.route('/patients/new', methods=['GET', 'POST'])
@login_required
def new_patient():
    insurances = Insurance.query.all()
    physicians = Physician.query.all()

    if request.method == 'POST':
        first_name = request.form.get('first_name')
        last_name = request.form.get('last_name')
        dob = request.form.get('dob')  # <-- get DOB here inside the request context
        dob_date = parse_dob(dob)      # <-- parse it here

        phone = request.form.get('phone')
        email = request.form.get('email')
        address = request.form.get('address')
        city = request.form.get('city')
        state = request.form.get('state')
        zip_code = request.form.get('zip_code')
        insurance_id = request.form.get('insurance') or None
        physician_id = request.form.get('physician') or None
        other_notes = request.form.get('other_notes')
        mrn = request.form.get('mrn')

        if not first_name or not last_name:
            flash("First name and last name are required.", "danger")
            return redirect(url_for('new_patient'))

        if dob and not dob_date:
            flash("Date of Birth must be in MM-DD-YYYY or MM/DD/YYYY format.", "danger")
            return redirect(url_for('new_patient'))

        patient = Patient(
            first_name=first_name,
            last_name=last_name,
            dob=dob_date,
            phone=phone,
            email=email,
            address=address,
            city=city,
            state=state,
            zip_code=zip_code,
            insurance_id=insurance_id,
            physician_id=physician_id,
            other_notes=other_notes,
            mrn=mrn
        )

        db.session.add(patient)
        db.session.commit()

        flash("New patient added!", "success")
        return redirect(url_for('patients_list'))

    # GET method: generate MRN for new patient form
    max_mrn_obj = db.session.query(func.max(Patient.mrn)).first()
    max_mrn = max_mrn_obj[0] if max_mrn_obj and max_mrn_obj[0] else None

    if max_mrn and max_mrn.isdigit():
        generated_mrn = str(int(max_mrn) + 1).zfill(7)
    else:
        generated_mrn = "1000001"

    return render_template(
        'add_patient.html',
        generated_mrn=generated_mrn,
        insurances=insurances,
        physicians=physicians
    )
    
# 3. Edit patient
@app.route('/patients/<int:patient_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_patient(patient_id):
    patient = Patient.query.get_or_404(patient_id)
    insurances = Insurance.query.all()
    physicians = Physician.query.all()
    if request.method == 'POST':
        patient.first_name = request.form.get('first_name')
        patient.last_name = request.form.get('last_name')
        dob = request.form.get('dob') or None
        patient.dob = datetime.strptime(dob, "%m-%d-%Y").date() if dob else None
        patient.phone = request.form.get('phone')
        patient.email = request.form.get('email')
        patient.address = request.form.get('address')
        insurance_id = request.form.get('insurance_id') or None
        physician_id = request.form.get('physician_id') or None
        patient.insurance_id = insurance_id if insurance_id else None
        patient.physician_id = physician_id if physician_id else None

        db.session.commit()
        flash("Patient updated!", "success")
        return redirect(url_for('patients_list'))
    return render_template('patient_form.html', patient=patient, insurances=insurances, physicians=physicians)
    
# ========== VISIT CRUD WITH GOOGLE CALENDAR SYNC ==========
@app.route('/visits')
@login_required
def visits_list():
    visits = Visit.query.order_by(Visit.visit_date.desc()).all()
    return render_template('visits_list.html', visits=visits)

@app.route('/visits/<int:visit_id>')
@login_required
def visit_detail(visit_id):
    visit = Visit.query.get_or_404(visit_id)
    notes = {}
    if visit.notes:
        try:
            notes = json.loads(visit.notes)
        except Exception:
            notes = {}

    # Clean up any old "goals" that were saved as a {"result": "..."} object
    goal_result = notes.get('goals', '')
    if isinstance(goal_result, str) and goal_result.strip().startswith('{'):
        try:
            goal_result = json.loads(goal_result).get('result', goal_result)
        except Exception:
            pass
        notes['goals'] = goal_result

    return render_template("visit_detail.html", visit=visit, notes=notes)

@app.route('/visits/new', methods=['GET', 'POST'])
@login_required
def new_visit():
    patients = Patient.query.all()
    therapists = Therapist.query.all()
    if request.method == 'POST':
        patient_id = request.form.get('patient_id')
        therapist_id = request.form.get('therapist_id')
        visit_type = request.form.get('visit_type')
        visit_date = request.form.get('visit_date')
        visit_time = request.form.get('visit_time')
        status = request.form.get('status', 'Scheduled')
        dt_start = datetime.strptime(f"{visit_date} {visit_time}", "%m-%d-%Y %H:%M")
        duration = 60
        visit = Visit(
            patient_id=patient_id,
            therapist_id=therapist_id,
            visit_type=visit_type,
            visit_date=dt_start,
            duration=duration,
            status=status
        )
        db.session.add(visit)
        db.session.commit()
        # Google Calendar sync
        event_id = create_google_event(visit)
        if event_id:
            visit.google_event_id = event_id
            db.session.commit()
        flash("Visit scheduled and synced!", "success")
        return redirect(url_for('visits_list'))
    return render_template('visit_form.html', patients=patients, therapists=therapists)

@app.route('/visits/<int:visit_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_visit(visit_id):
    visit = Visit.query.get_or_404(visit_id)
    notes = {}
    if visit.notes:
        try:
            notes = json.loads(visit.notes)
        except Exception:
            notes = {}

    if request.method == 'POST':
        # Update notes from submitted form fields
        notes['meddiag'] = request.form.get('meddiag', '').strip()
        notes['history'] = request.form.get('history', '').strip()
        notes['subjective'] = request.form.get('subjective', '').strip()
        notes['frequency'] = request.form.get('frequency', '').strip()
        notes['intervention'] = request.form.get('intervention', '').strip()
        notes['procedures'] = request.form.get('procedures', '').strip()
        # Add any other fields you want to save

        # Save JSON back to visit.notes
        visit.notes = json.dumps(notes)
        db.session.commit()

        # Sync Google Calendar event
        if visit.google_event_id:
            update_google_event(visit)
        else:
            event_id = create_google_event(visit)
            if event_id:
                visit.google_event_id = event_id
                db.session.commit()

        flash("Visit updated successfully!", "success")
        return redirect(url_for('visit_detail', visit_id=visit.id))

    # GET: render the form with current notes
    return render_template('edit_visit.html', visit=visit, notes=notes)

@app.route('/visits/<int:visit_id>/edit_note', methods=['GET', 'POST'])
@login_required
def edit_visit_note(visit_id):
    visit = Visit.query.get_or_404(visit_id)
    notes = {}
    if visit.notes:
        try:
            notes = json.loads(visit.notes)
        except Exception:
            notes = {}

    if request.method == 'POST':
        # Update note fields from form
        updated_notes = {}
        for key in notes.keys():
            updated_notes[key] = request.form.get(key, '').strip()
        # Update visit.notes JSON string
        visit.notes = json.dumps(updated_notes)
        db.session.commit()
        flash("Visit note updated!", "success")
        return redirect(url_for('visit_detail', visit_id=visit.id))

    return render_template('edit_visit_note.html', visit=visit, notes=notes)
    
@app.route('/visits/<int:visit_id>/delete', methods=['POST'])
@login_required
def delete_visit(visit_id):
    visit = Visit.query.get_or_404(visit_id)
    try:
        visit.status = 'Deleted'
        visit.deleted_at = datetime.utcnow()  # Set deleted timestamp
        db.session.commit()
        flash("Visit deleted successfully.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error deleting visit: {e}", "danger")
    return redirect(url_for('patient_detail', patient_id=visit.patient_id))

@app.route('/visits/<int:visit_id>/recover', methods=['POST'])
@login_required
def recover_visit(visit_id):
    visit = Visit.query.get_or_404(visit_id)
    try:
        visit.status = 'Scheduled'  # or previous status if stored
        visit.deleted_at = None
        db.session.commit()
        flash("Visit recovered successfully.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error recovering visit: {e}", "danger")
    return redirect(url_for('patient_detail', patient_id=visit.patient_id))


@app.route('/visits/<int:visit_id>/update_note', methods=['POST'])
@login_required
def update_visit_note(visit_id):
    visit = Visit.query.get_or_404(visit_id)
    notes = {}
    if visit.notes:
        try:
            notes = json.loads(visit.notes)
        except Exception:
            notes = {}

    # Update notes from form submission
    for key in notes.keys():
        notes[key] = request.form.get(key, '').strip()

    visit.notes = json.dumps(notes)
    db.session.commit()
    flash("Visit notes updated!", "success")
    # Redirect back to patient detail without editing mode
    return redirect(url_for('patient_detail', patient_id=visit.patient_id))

    
# ========== PHYSICIAN, INSURANCE, BILLING ==========

@app.route('/therapists/new', methods=['GET', 'POST'])
@login_required
def new_therapist():
    if request.method == 'POST':
        username = request.form.get('username')    # <--- ADD THIS
        password = request.form.get('password')    # <--- ADD THIS

        first_name = request.form.get('first_name')
        last_name = request.form.get('last_name')
        credentials = request.form.get('credentials')
        email = request.form.get('email')
        phone = request.form.get('phone')
        availability = request.form.get('availability')

        # Validate required fields
        if not username or not password:
            flash("Username and password are required.", "danger")
            return redirect(url_for('new_therapist'))

        if not first_name or not last_name or not email:
            flash("First name, last name, and email are required.", "danger")
            return redirect(url_for('new_therapist'))

        hashed_pw = generate_password_hash(password)

        therapist = Therapist(
            username=username,
            password=hashed_pw,
            first_name=first_name,
            last_name=last_name,
            credentials=credentials or "",
            email=email,
            phone=phone or "",
            availability=availability or ""
        )
        db.session.add(therapist)
        db.session.commit()
        flash("Therapist added!", "success")
        return redirect(url_for('therapists_list'))

    return render_template('therapist_form.html')
    
@app.route('/therapists')
@login_required
def therapists_list():
    therapists = Therapist.query.all()
    return render_template('therapists.html', therapists=therapists)
    
@app.route('/physicians')
@login_required
def physicians_list():
    physicians = Physician.query.all()
    return render_template('physicians.html', physicians=physicians)

@app.route('/insurance')
@login_required
def insurance_list():
    insurances = Insurance.query.all()
    return render_template('insurance_list.html', insurances=insurances)

@app.route('/billing')
@login_required
def billing_overview():
    return render_template('billing.html')

# ========== PT BUILDER & PATIENT PROFILE ==========
@app.route('/pt_builder')
@login_required
def pt_builder():
    return render_template('pt_builder.html')

@app.route('/patient/<int:patient_id>/daily_builder')
@login_required
def daily_builder(patient_id):
    patient = Patient.query.get_or_404(patient_id)
    return render_template('daily_builder.html', patient=patient)
    
@app.route('/patient/<int:patient_id>/pt_builder', methods=['GET'])
@login_required
def pt_builder_patient(patient_id):
    patient = Patient.query.get_or_404(patient_id)
    age = ""
    if patient.dob:
        try:
            if isinstance(patient.dob, str):
                birth = datetime.strptime(patient.dob, "%m-%d-%Y")
            else:
                birth = patient.dob
            today = datetime.today()
            age = today.year - birth.year - ((today.month, today.day) < (birth.month, birth.day))
        except Exception:
            age = ""
    return render_template('pt_builder_patient.html', patient=patient, age=age)
    
# ========== UPLOADS ==========
@app.route('/uploads', methods=['GET'])
@login_required
def uploads_list():
    attachments = Attachment.query.order_by(Attachment.uploaded_at.desc()).all()
    return render_template('uploads_list.html', attachments=attachments)

@app.route('/uploads/new', methods=['GET', 'POST'])
@login_required
def upload_file():
    patient_id = request.args.get('patient_id')  # Grab from URL if provided
    if request.method == 'POST':
        file = request.files.get('file')
        patient_id_form = request.form.get('patient_id') or patient_id
        visit_id = request.form.get('visit_id')
        if not file:
            flash("No file selected.", "danger")
            return redirect(url_for('upload_file', patient_id=patient_id))
        filename = file.filename
        save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(save_path)
        attachment = Attachment(
            filename=filename,
            filepath=save_path,
            patient_id=patient_id_form if patient_id_form else None,
            visit_id=visit_id if visit_id else None
        )
        db.session.add(attachment)
        db.session.commit()
        flash("File uploaded!", "success")
        # If you want to return to patient's profile after upload:
        if patient_id_form:
            return redirect(url_for('patient_profile', patient_id=patient_id_form))
        return redirect(url_for('uploads_list'))
    patients = Patient.query.all()
    visits = Visit.query.order_by(Visit.visit_date.desc()).all()
    return render_template('upload_form.html', patients=patients, visits=visits, selected_patient_id=patient_id)

@app.route('/uploads/<int:upload_id>/download')
@login_required
def download_file(upload_id):
    attachment = Attachment.query.get_or_404(upload_id)
    return send_file(attachment.filepath, as_attachment=True)

@app.route('/uploads/<int:upload_id>/delete', methods=['POST'])
@login_required
def delete_file(upload_id):
    attachment = Attachment.query.get_or_404(upload_id)
    try:
        if os.path.exists(attachment.filepath):
            os.remove(attachment.filepath)
        db.session.delete(attachment)
        db.session.commit()
        flash("File deleted.", "info")
    except Exception as e:
        flash(f"Error deleting file: {e}", "danger")
    # Redirect to patient profile if possible, else uploads list
    if attachment.patient_id:
        return redirect(url_for('patient_profile', patient_id=attachment.patient_id))
    return redirect(url_for('uploads_list'))
    
# ====== PT Section ======

PT_TEMPLATES = {
    "LBP Eval": {
        "meddiag": "",
        "history": "",
        "subjective": "Pt reports having LBP and is limiting daily functional activities. Pt would like to decrease pain and improve activity tolerance and return to PLOF. Pt agrees to PT evaluation.",
        "pain_location": "L-spine paraspinal, B QL, B gluteus medius",
        "pain_onset": "Chronic",
        "pain_condition": "Chronic",
        "pain_mechanism": "Muscle tension, stenosis, increased tone, structural changes",
        "pain_rating": "5/10, 0/10, 7/10",
        "pain_frequency": "Intermittent",
        "pain_description": "Sharp, Tense, Aching.",
        "pain_aggravating": "Sitting, standing, walking, forward bending, lifting/pulling.",
        "pain_relieved": "Pain meds prn and rest.",
        "pain_interferes": "Functional mobility, ADLs, sleep.",
        "meds": "See medication list",
        "tests": "N/A",
        "dme": "N/A",
        "plof": "Independent with mobility and ADLs",
        "posture": "Forward head lean, rounded shoulders, protracted scapular, slouch posture, decrease sitting postural awareness, loss of lumbar lordosis.",
        "rom": "Trunk Flexion: 50% limited\nTrunk Extension: 50% limited\nTrunk SB Left: 50% limited\nTrunk SB Right: 50% limited\nTrunk Rotation Left: 50% limited\nTrunk Rotation Right: 50% limited",
        "strength": "Gross Core Strength: 3/5\nGross Hip Strength: L/R  3/5; 3/5\nGross Knee Strength: L/R  3/5; 3/5\nGross Ankle Strength: L/R  3/5; 3/5",
        "palpation": "TTP: B QL, B gluteus medius, B piriformis, B paraspinal.\nJoint hypomobility: L1-L5 with central PA.\nIncreased paraspinal and gluteus medius tone",
        "functional": "Supine Sit Up Test: Unable\n30 seconds Chair Sit to Stand: 6x w/ increase LBP\nSingle Leg Balance Test: B LE: <1 sec with loss of balance.\nSingle Heel Raises Test: Unremarkable\nWalking on Toes:\nWalking on Heels:\nFunctional Squat:",
        "special": "(-) Slump Test\n(-) Unilateral SLR Test\n(-) Double SLR\n(-) Spring/Central PA\n(-) Piriformis test\n(-) SI Cluster Test",
        "impairments": "Prolonged sitting: 5 min\nStanding: 5 min\nWalking: 5 min\nBending, sweeping, cleaning, lifting: 5 min.",
        "goals": "Short-Term Goals (1–12 visits):\n1. Pt will report a reduction in low back pain to ≤1/10 to allow comfortable participation in functional activities.\n2. Pt will demonstrate a ≥10% improvement in trunk AROM to enhance mobility and reduce risk of reinjury during daily tasks.\n3. Pt will improve gross LE strength by at least 0.5 muscle grade to enhance safety during ADLs and minimize pain/injury risk.\n4. Pt will self-report ≥50% improvement in functional limitations related to ADLs.\nLong-Term Goals (13–25 visits):\n1. Pt will demonstrate B LE strength of ≥4/5 to independently and safely perform all ADLs.\n2. Pt will complete ≥14 repetitions on the 30-second chair sit-to-stand test to reduce fall risk.\n3. Pt will tolerate ≥30 minutes of activity to safely resume household tasks without limitation.\n4. Pt will demonstrate independence with HEP, using proper body mechanics and strength to support safe return to ADLs without difficulty.",
        "frequency": "1wk1, 2wk12",
        "intervention": "Manual Therapy (STM/IASTM/Joint Mob), Therapeutic Exercise, Therapeutic Activities, Neuromuscular Re-education, Gait Training, Balance Training, Pain Management Training, Modalities ice/heat 10-15min, E-Stim, Ultrasound, fall/injury prevention training, safety education/training, HEP education/training.",
        "procedures": "97161 Low Complexity\n97162 Moderate Complexity\n97163 High Complexity\n97140 Manual Therapy\n97110 Therapeutic Exercise\n97530 Therapeutic Activity\n97112 Neuromuscular Re-ed\n97116 Gait Training"
    },
    "Knee TKA Eval": {
        "meddiag": "",
        "history": "",
        "subjective": "Pt states s/p TKA and agreeable to PT evaluation. Pt reports having pain and swelling to the knee region and hasn't been using ice too much.",
        "pain_location": "Knee",
        "pain_onset": "",
        "pain_condition": "Acute",
        "pain_mechanism": "Post op swelling due to surgery",
        "pain_rating": "5/10, 3/10, 7/10",
        "pain_frequency": "Intermittent",
        "pain_description": "Sharp, Tension, Aching, dull/heaviness",
        "pain_aggravating": "Sitting, standing, walking, bed mobility.",
        "pain_relieved": "Pain meds prn, ice, rest, elevation",
        "pain_interferes": "Functional mobility, ADLs, sleep.",
        "meds": "See medication list",
        "tests": "N/A",
        "dme": "FWW",
        "plof": "Independent with mobility and ADLs.",
        "posture": "Forward head lean, rounded shoulders, protracted scapular, slouch posture, decrease sitting postural awareness, loss of lumbar lordosis.",
        "rom": "Hip Gross: WNL / WNL\nKnee Flex: \nKnee Ext:\nAnkle Gross: WNL / WNL",
        "strength": "Hip Gross: 4/5 / 4/5\nKnee Flex: 3/5* / 3/5*\nKnee Ext: 3/5* / 3/5*\nAnkle Gross: 4/5 / 4/5",
        "palpation": "TTP: B Quads, hamstring, knee swelling, warmth, tenderness periarticular",
        "functional": "Bed Mobility: SBA\n30 seconds Chair Sit to Stand: 2x w/ Knee pain\nSLB Test: Unable loss of balance\nSingle Heel Raises Test: 50% from full range, guarding at knee\nFunctional Squat: Unable",
        "special": "NT",
        "impairments": "Prolonged sitting: 5 min\nStanding: 5 min\nWalking: 5 min\nStep/stairs: 1 step",
        "goals": (
            "Short-Term Goals (1–12 visits):\n"
            "1. Pt will report a reduction in knee pain to ≤1/10 to allow safe and comfortable participation in functional activities.\n"
            "2. Pt will demonstrate a ≥10% improvement in knee AROM to enhance mobility and reduce risk of reinjury during daily tasks.\n"
            "3. Pt will improve gross LE strength by at least 0.5 muscle grade to enhance safety during ADLs and minimize pain/injury risk.\n"
            "4. Pt will self-report ≥50% improvement in functional limitations related to ADLs.\n"
            "Long-Term Goals (13–25 visits):\n"
            "1. Pt will demonstrate B LE strength of ≥4/5 to independently and safely perform all ADLs.\n"
            "2. Pt will complete ≥14 repetitions on the 30-second chair sit-to-stand test to reduce fall risk.\n"
            "3. Pt will tolerate ≥30 minutes of activity to safely resume household tasks without limitation.\n"
            "4. Pt will demonstrate independence with HEP, using proper body mechanics and strength to support safe return to ADLs without difficulty."
        ),
        "frequency": "1wk1, 2wk12",
        "intervention": "Manual Therapy (STM/IASTM/Joint Mob), Therapeutic Exercise, Therapeutic Activities, Neuromuscular Re-education, Gait Training, Balance Training, Pain Management Training, Modalities ice/heat 10-15min, E-Stim, Ultrasound, fall/injury prevention training, safety education/training, HEP education/training.",
        "procedures": "97161 Low Complexity\n97162 Moderate Complexity\n97163 High Complexity\n97140 Manual Therapy\n97110 Therapeutic Exercise\n97530 Therapeutic Activity\n97112 Neuromuscular Re-ed\n97116 Gait Training"
    }
}

def pt_parse_template(template):
    key_map = {
        "Medical Diagnosis": "meddiag",
        "Medical History/HNP": "history",
        "Subjective": "subjective",
        "Current Medication(s)": "meds",
        "Diagnostic Test(s)": "tests",
        "DME/Assistive Device": "dme",
        "PLOF": "plof",
        "Posture": "posture",
        "ROM": "rom",
        "Muscle Strength Test": "strength",
        "Palpation": "palpation",
        "Functional Test(s)": "functional",
        "Special Test(s)": "special",
        "Current Functional Mobility Impairment(s)": "impairments",
        "Goals": "goals",
        "Frequency/Duration": "frequency",
        "Intervention": "intervention",
        "Treatment Procedures": "procedures",
        "Area/Location of Injury": "pain_location",
        "Onset/Exacerbation Date": "pain_onset",
        "Condition of Injury": "pain_condition",
        "Mechanism of Injury": "pain_mechanism",
        "Pain Rating (P/B/W)": "pain_rating",
        "Pain Frequency": "pain_frequency",
        "Description": "pain_description",
        "Aggravating Factor": "pain_aggravating",
        "Relieved By": "pain_relieved",
        "Interferes With": "pain_interferes"
    }
    fields = {v: "" for v in key_map.values()}
    curr = None
    for line in template.splitlines():
        stripped = line.strip()
        matched = False
        for label, key in key_map.items():
            if stripped.startswith(label + ":"):
                curr = key
                _, val = stripped.split(":", 1)
                fields[key] = val.strip()
                matched = True
                break
        if not matched and curr and stripped:
            fields[curr] += "\n" + stripped
    return fields

@app.route("/pt_load_template", methods=["POST"])
@login_required
def pt_load_template():
    data = request.get_json()
    template_name = data.get("template", "")
    if not template_name:
        return jsonify(list(PT_TEMPLATES.keys()))
    else:
        return jsonify(PT_TEMPLATES.get(template_name, {}))


@app.route('/pt_export_word', methods=['POST'])
@login_required
def pt_export_word():
    data = request.json
    doc = pt_export_to_word(data)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='PT_Eval.docx'
    )

def pt_export_to_word(data):
    doc = Document()
    def add_separator():
        doc.add_paragraph('-' * 114)
    doc.add_paragraph(f"Medical Diagnosis: {data.get('meddiag', '')}")
    add_separator()
    doc.add_paragraph(f"Medical History/HNP:\n{data.get('history', '')}")
    add_separator()
    doc.add_paragraph(f"Subjective:\n{data.get('subjective', '')}")
    add_separator()
    doc.add_paragraph("Pain:")
    pain_fields = [
        ("Area/Location of Injury", "pain_location"),
        ("Onset/Exacerbation Date", "pain_onset"),
        ("Condition of Injury", "pain_condition"),
        ("Mechanism of Injury", "pain_mechanism"),
        ("Pain Rating (Present/Best/Worst)", "pain_rating"),
        ("Frequency", "pain_frequency"),
        ("Description", "pain_description"),
        ("Aggravating Factor", "pain_aggravating"),
        ("Relieved By", "pain_relieved"),
        ("Interferes With", "pain_interferes"),
    ]
    for label, key in pain_fields:
        doc.add_paragraph(f"{label}: {data.get(key, '')}")
    doc.add_paragraph(f"Current Medication(s): {data.get('meds', '')}")
    doc.add_paragraph(f"Diagnostic Test(s): {data.get('tests', '')}")
    doc.add_paragraph(f"DME/Assistive Device: {data.get('dme', '')}")
    doc.add_paragraph(f"PLOF: {data.get('plof', '')}")
    add_separator()
    doc.add_paragraph("Objective:")
    obj_fields = [
        ("Posture", "posture"),
        ("ROM", "rom"),
        ("Muscle Strength Test", "strength"),
        ("Palpation", "palpation"),
        ("Functional Test(s)", "functional"),
        ("Special Test(s)", "special"),
        ("Current Functional Mobility Impairment(s)", "impairments"),
    ]
    for label, key in obj_fields:
        doc.add_paragraph(f"{label}:")
        doc.add_paragraph(f"{data.get(key, '')}")
    add_separator()
    doc.add_paragraph("Assessment Summary:")
    doc.add_paragraph(data.get('summary', ''))
    add_separator()
    doc.add_paragraph("Goals:")
    doc.add_paragraph(data.get('goals', ''))
    add_separator()
    doc.add_paragraph("Frequency:")
    doc.add_paragraph(data.get('frequency', ''))
    add_separator()
    doc.add_paragraph("Intervention:")
    doc.add_paragraph(data.get('intervention', ''))
    add_separator()
    doc.add_paragraph("Treatment Procedures:")
    doc.add_paragraph(data.get('procedures', ''))
    add_separator()
    return doc

@app.route("/pt_export_pdf", methods=["POST"])
@login_required
def pt_export_pdf():
    data = request.get_json()
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 40

    def add_section(title, value):
        nonlocal y
        c.setFont("Helvetica-Bold", 13)
        c.drawString(40, y, title)
        y -= 18
        c.setFont("Helvetica", 11)
        for line in (value or "").split('\n'):
            c.drawString(48, y, line)
            y -= 14
            if y < 60:
                c.showPage()
                y = height - 40
        y -= 8
        c.setLineWidth(0.5)
        c.line(40, y, width - 40, y)
        y -= 16

    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, y, "Physical Therapy Evaluation")
    y -= 30

    add_section("Medical Diagnosis:", data.get("meddiag", ""))
    add_section("Medical History/HNP:", data.get("history", ""))
    add_section("Subjective:", data.get("subjective", ""))
    pain_lines = [
        f"Area/Location of Injury: {data.get('pain_location','')}",
        f"Onset/Exacerbation Date: {data.get('pain_onset','')}",
        f"Condition of Injury: {data.get('pain_condition','')}",
        f"Mechanism of Injury: {data.get('pain_mechanism','')}",
        f"Pain Rating (Present/Best/Worst): {data.get('pain_rating','')}",
        f"Frequency: {data.get('pain_frequency','')}",
        f"Description: {data.get('pain_description','')}",
        f"Aggravating Factor: {data.get('pain_aggravating','')}",
        f"Relieved By: {data.get('pain_relieved','')}",
        f"Interferes With: {data.get('pain_interferes','')}",
        "",
        f"Current Medication(s): {data.get('meds','')}",
        f"Diagnostic Test(s): {data.get('tests','')}",
        f"DME/Assistive Device: {data.get('dme','')}",
        f"PLOF: {data.get('plof','')}",
    ]
    add_section("Pain:", "\n".join(pain_lines))
    obj_lines = [
        f"Posture: {data.get('posture','')}",
        "",
        f"ROM: \n{data.get('rom','')}",
        "",
        f"Muscle Strength Test: \n{data.get('strength','')}",
        "",
        f"Palpation: \n{data.get('palpation','')}",
        "",
        f"Functional Test(s): \n{data.get('functional','')}",
        "",
        f"Special Test(s): \n{data.get('special','')}",
        "",
        f"Current Functional Mobility Impairment(s): \n{data.get('impairments','')}",
    ]
    add_section("Objective:", "\n".join(obj_lines))
    add_section("Assessment Summary:", data.get("summary", ""))
    add_section("Goals:", data.get("goals", ""))
    add_section("Frequency:", data.get("frequency", ""))
    add_section("Intervention:", data.get("intervention", ""))
    add_section("Treatment Procedures:", data.get("procedures", ""))

    c.save()
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name="PT_Eval.pdf",
        mimetype="application/pdf"
    )

# ====== OT Section ======

OT_TEMPLATES = {
    "OT Eval Template": """Medical Diagnosis:
Medical History/HNP:
Subjective: Pt reports upper extremity pain and is limiting ADLs. Pt would like to improve function and return to PLOF. Pt agrees to OT evaluation.
Pain:
Area/Location of Injury: R shoulder
Onset/Exacerbation Date: 3 weeks ago
Condition of Injury: Acute on chronic
Mechanism of Injury: Lifting
Pain Rating (P/B/W): 4/10, 1/10, 7/10
Pain Frequency: Intermittent
Description: Sharp, throbbing
Aggravating Factor: Overhead activity, reaching
Relieved By: Rest, ice
Interferes With: Grooming, dressing, bathing

Current Medication(s): See medication list

Diagnostic Test(s): MRI right shoulder

DME/Assistive Device: None

PLOF: Independent

Posture: Forward head, rounded shoulders

ROM: R shoulder flexion 100°, abduction 80°

Muscle Strength Test: R shoulder 3+/5

Palpation: TTP R supraspinatus

Functional Test(s): Unable to reach overhead

Special Test(s): (+) Impingement

Current Functional Mobility Impairment(s): Reaching, overhead activity

Goals:
Short-Term Goals (1–12 visits):
1. Pt will decrease pain to ≤2/10 during ADLs.
2. Pt will improve R shoulder ROM to 140° flexion.
3. Pt will improve strength to 4/5.
4. Pt will perform ADLs independently.

Long-Term Goals (13–25 visits):
1. Pt will maintain pain ≤1/10 with all activity.
2. Pt will achieve full ROM and strength in R shoulder.
3. Pt will return to all prior ADLs independently.
4. Pt will independently complete HEP.

Frequency/Duration: 2x/wk x 6wks

Intervention: Manual Therapy, TherEx, HEP training, ADL retraining

Treatment Procedures:
97165 OT Eval
97110 Ther Ex
97530 Ther Activity
97535 Self-care Mgmt
"""
}

def ot_parse_template(template):
    key_map = {
        "Medical Diagnosis": "ot_meddiag",
        "Medical History/HNP": "ot_history",
        "Subjective": "ot_subjective",
        "Current Medication(s)": "ot_meds",
        "Diagnostic Test(s)": "ot_tests",
        "DME/Assistive Device": "ot_dme",
        "PLOF": "ot_plof",
        "Posture": "ot_posture",
        "ROM": "ot_rom",
        "Muscle Strength Test": "ot_strength",
        "Palpation": "ot_palpation",
        "Functional Test(s)": "ot_functional",
        "Special Test(s)": "ot_special",
        "Current Functional Mobility Impairment(s)": "ot_impairments",
        "Goals": "ot_goals",
        "Frequency/Duration": "ot_frequency",
        "Intervention": "ot_intervention",
        "Treatment Procedures": "ot_procedures",
        "Area/Location of Injury": "ot_pain_location",
        "Onset/Exacerbation Date": "ot_pain_onset",
        "Condition of Injury": "ot_pain_condition",
        "Mechanism of Injury": "ot_pain_mechanism",
        "Pain Rating (P/B/W)": "ot_pain_rating",
        "Pain Frequency": "ot_pain_frequency",
        "Description": "ot_pain_description",
        "Aggravating Factor": "ot_pain_aggravating",
        "Relieved By": "ot_pain_relieved",
        "Interferes With": "ot_pain_interferes"
    }
    fields = {v: "" for v in key_map.values()}
    curr = None
    for line in template.splitlines():
        stripped = line.strip()
        matched = False
        for label, key in key_map.items():
            if stripped.startswith(label + ":"):
                curr = key
                _, val = stripped.split(":", 1)
                fields[key] = val.strip()
                matched = True
                break
        if not matched and curr and stripped:
            fields[curr] += "\n" + stripped
    return fields

@app.route("/ot_load_template", methods=["POST"])
@login_required
def ot_load_template():
    name = request.json.get("template", "")
    text = OT_TEMPLATES.get(name, "")
    return jsonify(ot_parse_template(text))

@app.route("/pt_generate_diffdx", methods=["POST"])
@login_required
def pt_generate_diffdx():
    f = request.json.get("fields", {})
    pain = "; ".join(f"{lbl}: {f.get(key,'')}"
                      for lbl,key in [
                          ("Area/Location", "pain_location"),
                          ("Onset", "pain_onset"),
                          ("Condition", "pain_condition"),
                          ("Mechanism", "pain_mechanism"),
                          ("Rating", "pain_rating"),
                          ("Frequency", "pain_frequency"),
                          ("Description", "pain_description"),
                          ("Aggravating", "pain_aggravating"),
                          ("Relieved", "pain_relieved"),
                          ("Interferes", "pain_interferes"),
                      ])
    prompt = (
        "You are a PT clinical assistant. Provide the single best-fit diagnosis Keep it clean.:\n\n"
        f"Subjective:\n{f.get('subjective','')}\n\n"
        f"Pain:\n{pain}\n\n"
        f"Objective:\nPosture: {f.get('posture','')}\n"
        f"ROM: {f.get('rom','')}\n"
        f"Strength: {f.get('strength','')}\n"
    )
    result = gpt_call(prompt, max_tokens=200)
    return jsonify({"result": result})


@app.route("/pt_generate_summary", methods=["POST"])
@login_required
def pt_generate_summary():
    f = request.json.get("fields", {})
    name = f.get("name", "Pt Name")
    age = f.get("age", "X")
    gender = f.get("gender", "patient").lower()
    pmh = f.get("history", "no significant history")
    today = f.get("currentdate", date.today().strftime("%m/%d/%Y"))
    subj = f.get("subjective", "")
    moi = f.get("pain_mechanism", "")
    dx = f.get("diffdx", "")
    strg = f.get("strength", "")
    rom = f.get("rom", "")
    impair = f.get("impairments", "")
    func = f.get("functional", "")

    prompt = (
        "Generate a concise, 7-8 sentence Physical Therapy assessment summary medicare compliant for PT documentation. "
        "Use clinical, professional language and use abbreviations only (e.g., HEP, ADLs, LBP, STM, TherEx, etc.; "
        "do not spell out the abbreviation and do not write both full term and abbreviation). "
        "Never use the phrase 'The patient'; instead, use 'Pt' at the start of each relevant sentence. "
        "Make sure to the summary is well written with no gramatical error and health care professional manner."
        f"Start with: \"{name}, a {age} y/o {gender} with relevant history of {pmh}.\" "
        f"Include: "
        f"How/when/why pt was seen (PT initial eval on {today} for {subj}), "
        f"mechanism of injury if available ({moi}), "
        f"main differential dx ({dx}), "
        f"current impairments Summary, but not too specific(strength: {strg}; ROM: {rom}; balance/mobility: {impair}), "
        f"functional/activity/participation limitations: {func}, "
        "a professional prognosis and that skilled PT will help pt address impairments and return to PLOF. "
        "Do not use bulleted or numbered lists—just a single, well-written summary paragraph."
    )
    result = gpt_call(prompt, max_tokens=350)
    return jsonify({"result": result})


@app.route('/pt_generate_goals', methods=['POST'])
@login_required
def pt_generate_goals():
    fields = request.json.get("fields", {})
    prompt = """
    You are a clinical assistant helping a PT write documentation.
    Using ONLY the provided eval info (summary, objective findings, strength, ROM, impairments, and functional limitations),
    generate clinically-appropriate, Medicare-compliant short-term and long-term PT goals.
    ALWAYS follow this exact format—do not add, skip, reorder, or alter any lines or labels.
    DO NOT add any explanations, introductions, dashes, bullets, or extra indentation. Output ONLY this structure:

    Short-Term Goals (1–12 visits):
    1. [goal statement]
    2. [goal statement]
    3. [goal statement]
    4. [goal statement]

    Long-Term Goals (13–25 visits):
    1. [goal statement]
    2. [goal statement]
    3. [goal statement]
    4. [goal statement]
    """

    result = gpt_call(prompt, max_tokens=350)
    return jsonify({"result": result})
    


@app.route('/pt_generate_daily_summary', methods=['POST'])
@login_required
def pt_generate_daily_summary():
    data = request.json
    prompt = (
        "You are a physical therapist. "
        "Write a 6-sentence daily PT note summary in paragraph form. "
        "Use professional tone, refer to 'patient' (not 'the patient' or 'patient reported'). "
        "Summarize the following:\n"
        f"Diagnosis: {data.get('diagnosis','')}\n"
        f"Interventions: {data.get('interventions','')}\n"
        f"Tx Tolerance: {data.get('tolerance','')}\n"
        f"Current Progress: {data.get('progress','')}\n"
        f"Next Visit Plan: {data.get('plan','')}\n"
        "Do not use the phrases 'patient reported' or 'the patient'. "
        "Do not spell out, use abbreviation only, avoid using both next to each other. "
        "After summarizes skip a row write a 1-2 sentences for next visit plan of care utilizing something along Focusing on PT POC to improve strength, endurance, mechanics, activity tolerance with manual therapy, ther-ex, ther-act, IASTM. Improve activity tolerance to return to safe ADLs and community participation and ambulation."
    )
    try:
        completion = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=250
        )
        summary = completion.choices[0].message.content.strip()
        return jsonify({"result": summary})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
        
@app.route('/ot_export_word', methods=['POST'])
@login_required
def ot_export_word():
    data = request.json
    doc = ot_export_to_word(data)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='OT_Eval.docx'
    )

def ot_export_to_word(data):
    doc = Document()
    def add_separator():
        doc.add_paragraph('-' * 114)
    doc.add_paragraph(f"Medical Diagnosis: {data.get('ot_meddiag', '')}")
    add_separator()
    doc.add_paragraph(f"Medical History/HNP:\n{data.get('ot_history', '')}")
    add_separator()
    doc.add_paragraph(f"Subjective:\n{data.get('ot_subjective', '')}")
    add_separator()
    doc.add_paragraph("Pain:")
    pain_fields = [
        ("Area/Location of Injury", "ot_pain_location"),
        ("Onset/Exacerbation Date", "ot_pain_onset"),
        ("Condition of Injury", "ot_pain_condition"),
        ("Mechanism of Injury", "ot_pain_mechanism"),
        ("Pain Rating (Present/Best/Worst)", "ot_pain_rating"),
        ("Frequency", "ot_pain_frequency"),
        ("Description", "ot_pain_description"),
        ("Aggravating Factor", "ot_pain_aggravating"),
        ("Relieved By", "ot_pain_relieved"),
        ("Interferes With", "ot_pain_interferes"),
    ]
    for label, key in pain_fields:
        doc.add_paragraph(f"{label}: {data.get(key, '')}")
    doc.add_paragraph(f"Current Medication(s): {data.get('ot_meds', '')}")
    doc.add_paragraph(f"Diagnostic Test(s): {data.get('ot_tests', '')}")
    doc.add_paragraph(f"DME/Assistive Device: {data.get('ot_dme', '')}")
    doc.add_paragraph(f"PLOF: {data.get('ot_plof', '')}")
    add_separator()
    doc.add_paragraph("Objective:")
    obj_fields = [
        ("Posture", "ot_posture"),
        ("ROM", "ot_rom"),
        ("Muscle Strength Test", "ot_strength"),
        ("Palpation", "ot_palpation"),
        ("Functional Test(s)", "ot_functional"),
        ("Special Test(s)", "ot_special"),
        ("Current Functional Mobility Impairment(s)", "ot_impairments"),
    ]
    for label, key in obj_fields:
        doc.add_paragraph(f"{label}:")
        doc.add_paragraph(f"{data.get(key, '')}")
    add_separator()
    doc.add_paragraph("Assessment Summary:")
    doc.add_paragraph(data.get('ot_summary', ''))
    add_separator()
    doc.add_paragraph("Goals:")
    doc.add_paragraph(data.get('ot_goals', ''))
    add_separator()
    doc.add_paragraph("Frequency:")
    doc.add_paragraph(data.get('ot_frequency', ''))
    add_separator()
    doc.add_paragraph("Intervention:")
    doc.add_paragraph(data.get('ot_intervention', ''))
    add_separator()
    doc.add_paragraph("Treatment Procedures:")
    doc.add_paragraph(data.get('ot_procedures', ''))
    add_separator()
    return doc

@app.route("/ot_export_pdf", methods=["POST"])
@login_required
def ot_export_pdf():
    data = request.get_json()
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 40

    def add_section(title, value):
        nonlocal y
        c.setFont("Helvetica-Bold", 13)
        c.drawString(40, y, title)
        y -= 18
        c.setFont("Helvetica", 11)
        for line in (value or "").split('\n'):
            c.drawString(48, y, line)
            y -= 14
            if y < 60:
                c.showPage()
                y = height - 40
        y -= 8
        c.setLineWidth(0.5)
        c.line(40, y, width - 40, y)
        y -= 16

    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, y, "Occupational Therapy Evaluation")
    y -= 30

    add_section("Medical Diagnosis:", data.get("ot_meddiag", ""))
    add_section("Medical History/HNP:", data.get("ot_history", ""))
    add_section("Subjective:", data.get("ot_subjective", ""))
    pain_lines = [
        f"Area/Location of Injury: {data.get('ot_pain_location','')}",
        f"Onset/Exacerbation Date: {data.get('ot_pain_onset','')}",
        f"Condition of Injury: {data.get('ot_pain_condition','')}",
        f"Mechanism of Injury: {data.get('ot_pain_mechanism','')}",
        f"Pain Rating (Present/Best/Worst): {data.get('ot_pain_rating','')}",
        f"Frequency: {data.get('ot_pain_frequency','')}",
        f"Description: {data.get('ot_pain_description','')}",
        f"Aggravating Factor: {data.get('ot_pain_aggravating','')}",
        f"Relieved By: {data.get('ot_pain_relieved','')}",
        f"Interferes With: {data.get('ot_pain_interferes','')}",
        "",
        f"Current Medication(s): {data.get('ot_meds','')}",
        f"Diagnostic Test(s): {data.get('ot_tests','')}",
        f"DME/Assistive Device: {data.get('ot_dme','')}",
        f"PLOF: {data.get('ot_plof','')}",
    ]
    add_section("Pain:", "\n".join(pain_lines))
    obj_lines = [
        f"Posture: {data.get('ot_posture','')}",
        "",
        f"ROM: \n{data.get('ot_rom','')}",
        "",
        f"Muscle Strength Test: \n{data.get('ot_strength','')}",
        "",
        f"Palpation: \n{data.get('ot_palpation','')}",
        "",
        f"Functional Test(s): \n{data.get('ot_functional','')}",
        "",
        f"Special Test(s): \n{data.get('ot_special','')}",
        "",
        f"Current Functional Mobility Impairment(s): \n{data.get('ot_impairments','')}",
    ]
    add_section("Objective:", "\n".join(obj_lines))
    add_section("Assessment Summary:", data.get("ot_summary", ""))
    add_section("Goals:", data.get("ot_goals", ""))
    add_section("Frequency:", data.get("ot_frequency", ""))
    add_section("Intervention:", data.get("ot_intervention", ""))
    add_section("Treatment Procedures:", data.get("ot_procedures", ""))

    c.save()
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name="OT_Eval.pdf",
        mimetype="application/pdf"
    )
    
# ========== GPT HELPER ==========

def gpt_call(prompt, max_tokens=350):
    if not client:
        return "OpenAI client not configured"
    try:
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=max_tokens
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        return f"OpenAI error: {e}"


# ========== MAIN ==========

if __name__ == "__main__":
    with app.app_context():
        db.create_all()
    app.run(debug=True)

